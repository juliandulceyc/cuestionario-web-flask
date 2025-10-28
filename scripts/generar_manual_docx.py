"""
Script para generar un archivo Word (.docx) del manual de usuario.

Requisitos: python-docx (instalado desde requirements.txt)
Ejecutar desde la raíz del proyecto:

    python scripts/generar_manual_docx.py

Salida: docs/Manual_de_Usuario.docx
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os

MANUAL_MD_PATH = os.path.join('docs', 'Manual_de_Usuario.md')
OUTPUT_DOCX_PATH = os.path.join('docs', 'Manual_de_Usuario.docx')


def add_heading(document: Document, text: str, level: int = 1):
    h = document.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph(document: Document, text: str):
    p = document.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)


def md_to_docx(document: Document, md_lines):
    """Conversión simple de Markdown a docx:
    - Encabezados #, ##, ### -> heading niveles 1,2,3
    - Listas con '-' -> viñetas
    - Resto -> párrafos
    """
    for raw in md_lines:
        line = raw.rstrip('\n')
        if not line.strip():
            document.add_paragraph('')
            continue
        if line.startswith('### '):
            add_heading(document, line[4:].strip(), level=3)
        elif line.startswith('## '):
            add_heading(document, line[3:].strip(), level=2)
        elif line.startswith('# '):
            add_heading(document, line[2:].strip(), level=1)
        elif line.lstrip().startswith('- '):
            document.add_paragraph(line.lstrip()[2:].strip(), style='List Bullet')
        else:
            add_paragraph(document, line)


def main():
    os.makedirs('docs', exist_ok=True)
    doc = Document()

    # Portada
    add_heading(doc, 'Manual de Usuario — Sistema de Evaluación Técnica', level=1)
    add_paragraph(doc, f'Fecha: {datetime.now().strftime("%d/%m/%Y")}')
    add_paragraph(doc, 'Versión del sistema: 1.0')
    doc.add_page_break()

    # Tabla de contenido (TOC)
    add_heading(doc, 'Tabla de contenido', level=1)
    # Inserta un campo TOC (Word) para niveles 1-3 con hipervínculos
    p = doc.add_paragraph()
    r = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    r._r.append(fld_separate)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r2 = p.add_run()
    r2._r.append(fld_end)

    add_paragraph(doc, ' ')  # espacio
    doc.add_page_break()

    # Contenido desde el Markdown
    if os.path.exists(MANUAL_MD_PATH):
        with open(MANUAL_MD_PATH, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        md_to_docx(doc, lines)
    else:
        add_paragraph(doc, 'No se encontró docs/Manual_de_Usuario.md')

    doc.save(OUTPUT_DOCX_PATH)
    print(f'Documento Word generado en: {OUTPUT_DOCX_PATH}')


if __name__ == '__main__':
    main()
